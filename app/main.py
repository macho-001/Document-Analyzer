from agents.graph import run_agent
import json
import os
from pathlib import Path


def load_document(file_path: str) -> dict:
    """
    Load and parse document from file.
    
    Supports: PDF, DOCX, TXT, MD
    
    Args:
        file_path: Path to document file
        
    Returns:
        Dict containing document content and metadata
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Document not found: {file_path}")
    
    file_path = Path(file_path)
    file_ext = file_path.suffix.lower()
    
    print(f"\n📄 Loading document: {file_path.name}")
    print(f"📝 File type: {file_ext}")
    
    try:
        if file_ext == '.pdf':
            return load_pdf(file_path)
        elif file_ext == '.docx':
            return load_docx(file_path)
        elif file_ext in ['.txt', '.md']:
            return load_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
            
    except Exception as e:
        print(f"❌ Error loading document: {str(e)}")
        raise

def load_pdf(file_path: Path) -> dict:
    """Load PDF using pdfplumber for better image and diagram detection."""
    try:
        import pdfplumber
        
        full_text = []
        num_images = 0
        num_tables = 0
        has_vector_graphics = False
        
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                # Extract text
                text = page.extract_text()
                if text:
                    full_text.append(text)
                
                # Count actual image objects
                num_images += len(page.images)
                
                # Detect tables using pdfplumber's built-in table finder
                tables = page.find_tables()
                num_tables += len(tables)
                
                # Logic for vector diagrams: if a page has many lines/rects 
                # but few images, it's likely a vector diagram or chart
                if len(page.lines) > 10 or len(page.rects) > 10:
                    has_vector_graphics = True
        
        content = '\n'.join(full_text)
        sections = extract_sections_from_text(content)
        
        return {
            'content': content,
            'file_path': str(file_path),
            'file_type': 'pdf',
            'metadata': {
                'filename': file_path.name,
                'num_pages': len(pdf.pages),
                'num_images': num_images,
                'num_tables': num_tables,
                'has_vector_graphics': has_vector_graphics, # Unique to PDF
                'sections': sections,
                'file_size': os.path.getsize(file_path)
            }
        }
    except ImportError:
        print("⚠️ pdfplumber not installed. Install with: pip install pdfplumber")
        raise
    except Exception as e:
        print(f"❌ Error reading PDF: {str(e)}")
        raise
# def load_pdf(file_path: Path) -> dict:
#     """Load PDF document using PyPDF2 or pdfplumber."""
#     try:
#         import PyPDF2
        
#         with open(file_path, 'rb') as file:
#             pdf_reader = PyPDF2.PdfReader(file)
#             num_pages = len(pdf_reader.pages)
            
#             # Extract text from all pages
#             content = []
#             for page_num in range(num_pages):
#                 page = pdf_reader.pages[page_num]
#                 content.append(page.extract_text())
            
#             full_text = '\n'.join(content)
            
#             # Try to extract sections/headings
#             sections = extract_sections_from_text(full_text)
            
#             return {
#                 'content': full_text,
#                 'file_path': str(file_path),
#                 'file_type': 'pdf',
#                 'metadata': {
#                     'filename': file_path.name,
#                     'num_pages': num_pages,
#                     'sections': sections,
#                     'file_size': os.path.getsize(file_path)
#                 }
#             }
#     except ImportError:
#         print("⚠️  PyPDF2 not installed. Install with: pip install PyPDF2")
#         raise
#     except Exception as e:
#         print(f"❌ Error reading PDF: {str(e)}")
#         raise


def load_docx(file_path: Path) -> dict:
    """Load DOCX document using python-docx."""
    try:
        from docx import Document
        
        doc = Document(file_path)
        
        # Extract text from paragraphs
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        full_text = '\n'.join(paragraphs)
        
        # Extract headings
        sections = []
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                sections.append(para.text)
        
        # Count tables and images
        num_tables = len(doc.tables)
        num_images = len([r for r in doc.inline_shapes])
        
        return {
            'content': full_text,
            'file_path': str(file_path),
            'file_type': 'docx',
            'metadata': {
                'filename': file_path.name,
                'num_paragraphs': len(paragraphs),
                'sections': sections,
                'num_tables': num_tables,
                'num_images': num_images,
                'file_size': os.path.getsize(file_path)
            }
        }
    except ImportError:
        print("⚠️  python-docx not installed. Install with: pip install python-docx")
        raise
    except Exception as e:
        print(f"❌ Error reading DOCX: {str(e)}")
        raise


def load_text(file_path: Path) -> dict:
    """Load plain text or markdown file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        sections = extract_sections_from_text(content)
        
        return {
            'content': content,
            'file_path': str(file_path),
            'file_type': file_path.suffix[1:],  # txt or md
            'metadata': {
                'filename': file_path.name,
                'num_lines': len(content.split('\n')),
                'sections': sections,
                'file_size': os.path.getsize(file_path)
            }
        }
    except Exception as e:
        print(f"❌ Error reading text file: {str(e)}")
        raise


def extract_sections_from_text(text: str) -> list:
    """
    Extract section headings from text.
    Looks for common heading patterns.
    """
    import re
    
    sections = []
    
    # Pattern 1: Lines that are all caps
    for line in text.split('\n'):
        line = line.strip()
        if line and line.isupper() and len(line) > 3 and len(line) < 100:
            sections.append(line)
    
    # Pattern 2: Markdown headers (# Header)
    md_headers = re.findall(r'^#+\s+(.+)$', text, re.MULTILINE)
    sections.extend(md_headers)
    
    # Pattern 3: Common section keywords
    keywords = ['overview', 'introduction', 'abstract', 'summary', 
                'methodology', 'results', 'conclusion', 'references']
    for line in text.split('\n'):
        line_lower = line.strip().lower()
        if any(keyword in line_lower for keyword in keywords):
            if len(line.strip()) < 100:  # Likely a heading
                sections.append(line.strip())
    
    # Remove duplicates while preserving order
    seen = set()
    unique_sections = []
    for section in sections:
        if section.lower() not in seen:
            seen.add(section.lower())
            unique_sections.append(section)
    
    return unique_sections[:20]  # Limit to first 20 sections


def main():
    """Main entry point for autonomous agent."""
    print("="*70)
    print("🤖 AUTONOMOUS DOCUMENT ANALYSIS AGENT")
    print("="*70)
    
    # Get document path
    document_path = input("\n📄 Enter document path: ").strip()
    
    if not document_path:
        print("❌ No document path provided. Exiting.")
        return
    
    # Load document
    try:
        document = load_document(document_path)
        print(f"✅ Document loaded successfully!")
        print(f"   Pages/Lines: {document['metadata'].get('num_pages') or document['metadata'].get('num_lines')}")
        print(f"   Sections found: {len(document['metadata']['sections'])}")
        if document['metadata']['sections']:
            print(f"   Sample sections: {', '.join(document['metadata']['sections'][:3])}")
    except Exception as e:
        print(f"❌ Failed to load document: {str(e)}")
        return
    
    # Get user query
    print("\n" + "="*70)
    query = input("💬 Enter your question about the document: ").strip()
    
    if not query:
        print("❌ No query provided. Exiting.")
        return
    
    # Run the agent
    final_state = run_agent(
        query=query,
        document=document,
        verbose=True
    )
    
    # Save results
    save_results = input("\n💾 Save results to file? (y/n): ").strip().lower()
    if save_results == 'y':
        output_file = 'agent_results.json'
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump({
                'query': final_state['query'],
                'document_file': document['file_path'],
                'goal': final_state['goal'],
                'plan': final_state['plan'],
                'observations': final_state['observations'],
                'tool_outputs': final_state['tool_outputs'],
                'final_answer': final_state['final_answer'],
                'status': final_state['status']
            }, f, indent=2, ensure_ascii=False)
        print(f"✅ Results saved to {output_file}")


if __name__ == "__main__":
    main()